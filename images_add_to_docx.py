from docx import Document
import os

folder_path = input("folder_path: ")
aligment = input("aligment, 0:left, 1:centre (default), 2: right\ncode : ")
aligment = "1" if aligment == "" else aligment
aligment = int(0 if aligment not in "012" else aligment)
print("\nImages:")

# Retrieve the folder name
folder_name = os.path.basename(folder_path)

# Create a new Word document
doc = Document()

# Get a list of PNG files in the folder and sort them by creation time
png_files = sorted(
    [file for file in os.listdir(folder_path) if file.lower().endswith('.png')],
    key=lambda f: os.path.getmtime(os.path.join(folder_path, f)),
    reverse=False
)


# Add each PNG file name to the Word document
for png_file in png_files:
    # Remove the file extension (.png) from the file name
    file_name = os.path.splitext(png_file)[0]

    # Add a newline 
    paragraph = doc.add_paragraph()
    paragraph.alignment = 0
        
    # Set the picture alignment to center
    paragraph = doc.add_paragraph()
    paragraph.alignment = aligment  # alignment code
    run = paragraph.add_run()
    run.add_picture(os.path.join(folder_path, png_file))
    print(png_file)

    # Add the file name below the image
    paragraph = doc.add_paragraph()
    paragraph.alignment = 1  # Center alignment
    paragraph.add_run(file_name)

# Save the Word document with the folder name in the input folder
doc.save(os.path.join(folder_path, folder_name + '.docx'))
print("\nSuccessful add to : " + os.path.join(folder_path, folder_name + '.docx'))
input()